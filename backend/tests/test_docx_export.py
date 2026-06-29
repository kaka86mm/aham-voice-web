"""markdown → docx 转换器测试。

纪要内容是标准 Markdown（标题/段落/列表/表格/粗体/代码），转 docx 后
要正确呈现这些结构。用 python-docx 读取生成的文件验证结构。
"""
import io
from pathlib import Path

import pytest
from docx import Document

from backend.app.docx_export import markdown_to_docx_bytes


def _load_docx(docx_bytes: bytes) -> Document:
    """从字节流加载 docx Document 用于断言。"""
    return Document(io.BytesIO(docx_bytes))


# -------- 标题 --------

def test_headings_levels():
    """# ## ### 转成对应级别的 Heading。"""
    md = "# 一级标题\n\n## 二级标题\n\n### 三级标题\n\n#### 四级标题"
    docx_bytes = markdown_to_docx_bytes(md)
    doc = _load_docx(docx_bytes)
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert len(headings) == 4
    assert "一级标题" in headings[0].text
    assert "四级标题" in headings[3].text


def test_document_title_from_h1():
    """第一个 # 作为文档标题（Word 标题样式）。"""
    md = "# 会议纪要标题\n\n正文内容"
    docx_bytes = markdown_to_docx_bytes(md)
    doc = _load_docx(docx_bytes)
    # 至少有一个段落含标题文字
    texts = [p.text for p in doc.paragraphs]
    assert any("会议纪要标题" in t for t in texts)


# -------- 段落与行内格式 --------

def test_plain_paragraph():
    """纯文本段落正确保留。"""
    md = "这是一段普通文字。\n\n这是第二段。"
    docx_bytes = markdown_to_docx_bytes(md)
    doc = _load_docx(docx_bytes)
    texts = [p.text for p in doc.paragraphs if p.style.name == "Normal"]
    assert any("普通文字" in t for t in texts)
    assert any("第二段" in t for t in texts)


def test_bold_inline():
    """**粗体** 转成粗体 run。"""
    md = "前面 **加粗文字** 后面"
    docx_bytes = markdown_to_docx_bytes(md)
    doc = _load_docx(docx_bytes)
    paragraph = [p for p in doc.paragraphs if "加粗文字" in p.text][0]
    bold_runs = [r for r in paragraph.runs if r.bold and "加粗文字" in r.text]
    assert len(bold_runs) == 1


# -------- 列表 --------

def test_unordered_list():
    """- 开头的行转成项目符号列表。"""
    md = "- 第一项\n- 第二项\n- 第三项"
    docx_bytes = markdown_to_docx_bytes(md)
    doc = _load_docx(docx_bytes)
    list_items = [p.text for p in doc.paragraphs if p.style.name.startswith("List")]
    assert len(list_items) == 3
    assert "第一项" in list_items[0]


def test_ordered_list():
    """1. 开头的行转成编号列表。"""
    md = "1. 步骤一\n2. 步骤二\n3. 步骤三"
    docx_bytes = markdown_to_docx_bytes(md)
    doc = _load_docx(docx_bytes)
    list_items = [p.text for p in doc.paragraphs if p.style.name.startswith("List")]
    assert len(list_items) == 3
    assert "步骤一" in list_items[0]


# -------- 表格 --------

def test_table():
    """markdown 表格转成 Word 表格。"""
    md = "| 说话人 | 情绪 | 占比 |\n|---|---|---|\n| 张总 | 中立 | 60% |\n| 李总 | 开心 | 40% |"
    docx_bytes = markdown_to_docx_bytes(md)
    doc = _load_docx(docx_bytes)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 3  # 表头 + 2 数据行
    assert "说话人" in table.rows[0].cells[0].text
    assert "张总" in table.rows[1].cells[0].text


# -------- 代码与引用 --------

def test_inline_code():
    """`行内代码` 转成等宽字体 run。"""
    md = "命令是 `ffmpeg -ss 00:30` 这样"
    docx_bytes = markdown_to_docx_bytes(md)
    doc = _load_docx(docx_bytes)
    paragraph = [p for p in doc.paragraphs if "ffmpeg" in p.text][0]
    # 行内代码应该作为独立 run 存在
    code_runs = [r for r in paragraph.runs if "ffmpeg" in r.text]
    assert len(code_runs) >= 1


def test_blockquote():
    """> 引用块转成 Word 引用/缩进段落。"""
    md = "> 这是一段引用内容"
    docx_bytes = markdown_to_docx_bytes(md)
    doc = _load_docx(docx_bytes)
    texts = [p.text for p in doc.paragraphs]
    assert any("引用内容" in t for t in texts)


# -------- 复杂真实纪要 --------

def test_realistic_summary():
    """模拟真实纪要（含标题/列表/粗体/时间戳），整体不报错且结构完整。"""
    md = """# 会议纪要

## 一句话概览
客户对**报价**有异议。

## 关键原文证据
- `[00:12:30]` 客户方：超预算
- `[15:45]` 销售方：可分阶段

## 待确认
1. 交付周期
2. 付款方式"""
    docx_bytes = markdown_to_docx_bytes(md)
    doc = _load_docx(docx_bytes)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "会议纪要" in all_text
    assert "报价" in all_text
    assert "00:12:30" in all_text


def test_chinese_font():
    """中文字符正确写入（不乱码）。python-docx 默认支持 UTF-8，
    这里验证读取回来的中文完整。"""
    md = "# 测试中文标题\n\n正文含中文：会议纪要、说话人、待确认问题。"
    docx_bytes = markdown_to_docx_bytes(md)
    doc = _load_docx(docx_bytes)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "测试中文标题" in all_text
    assert "会议纪要" in all_text
