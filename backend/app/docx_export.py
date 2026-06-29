"""会议纪要 Markdown → docx 转换。

国内会议纪要主流是 docx，比 md 更适合给非技术人员看。这里用 python-docx
自己构建 Word 文档，避免引入 pandoc/libreoffice 等重型系统依赖。

流程：markdown 文本 → markdown 库解析成 HTML → HTMLParser 遍历 → python-docx 构建。
覆盖纪要常用的结构：标题、段落、列表（有序/无序）、表格、粗体、行内代码、
引用块。行内格式（<strong>/<em>/<code>）通过跟踪标签状态，在 handle_data 时
创建带格式的 run。
"""
from __future__ import annotations

import html as html_module
from html.parser import HTMLParser
from io import BytesIO
from typing import Any

import markdown as md_lib
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def _decode_entities(text: str) -> str:
    """HTML 实体解码（&amp; &lt; &quot; 等）。"""
    return html_module.unescape(text)


# ──────────────────────────────────────────────────────────────────────────
# 行内片段：记录一段文本及其格式标记
# ──────────────────────────────────────────────────────────────────────────
class _InlineSpan:
    """一段连续的文本 + 它的格式标记（粗体/斜体/代码）。"""

    __slots__ = ("text", "bold", "italic", "code")

    def __init__(self, text: str, bold: bool = False, italic: bool = False, code: bool = False) -> None:
        self.text = text
        self.bold = bold
        self.italic = italic
        self.code = code


def _spans_to_runs(paragraph: Any, spans: list[_InlineSpan]) -> None:
    """把行内片段列表写入 paragraph 的 runs，应用对应格式。"""
    for span in spans:
        text = span.text
        if not text:
            continue
        run = paragraph.add_run(text)
        if span.bold:
            run.bold = True
        if span.italic:
            run.italic = True
        if span.code:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(10)


# ──────────────────────────────────────────────────────────────────────────
# HTML → docx 解析器
# ──────────────────────────────────────────────────────────────────────────
class _DocxBuilder(HTMLParser):
    """遍历 HTML 树，用 python-docx 构建 Word 文档。

    行内格式通过跟踪 <strong>/<em>/<code> 标签状态：遇到这些标签时设标志，
    handle_data 收集的文本就带上对应格式，存入 _current_spans。
    遇到块级元素结束（</p>/<li>/<h>）时，把 _current_spans flush 成段落。
    """

    def __init__(self, document: Document) -> None:
        super().__init__(convert_charrefs=True)
        self.doc = document
        # 行内格式状态
        self._bold = False
        self._italic = False
        self._code = False
        self._current_spans: list[_InlineSpan] = []
        # 块级上下文
        self._list_depth = 0
        self._in_ordered = False
        # 表格
        self._in_table = False
        self._table_rows: list[list[list[_InlineSpan]]] = []
        self._current_row: list[list[_InlineSpan]] = []
        self._current_cell: list[_InlineSpan] = []
        # 引用块缩进
        self._blockquote_depth = 0

    def _flush_inline_to_span(self, text: str) -> None:
        """把当前文本（带当前格式状态）存为一个 span。"""
        if not text:
            return
        span = _InlineSpan(text, bold=self._bold, italic=self._italic, code=self._code)
        # 表格场景：存到 cell
        if self._in_table and self._current_cell is not None and not isinstance(self._current_cell, bool):
            self._current_cell.append(span)
        # 普通段落/标题/列表项场景：存到 _current_spans
        else:
            self._current_spans.append(span)

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._heading_level = int(tag[1])
        elif tag == "p":
            self._current_spans = []
        elif tag == "ul":
            self._list_depth += 1
            self._in_ordered = False
        elif tag == "ol":
            self._list_depth += 1
            self._in_ordered = True
        elif tag == "li":
            self._current_spans = []
        elif tag == "blockquote":
            self._blockquote_depth += 1
        elif tag == "strong" or tag == "b":
            self._bold = True
        elif tag == "em" or tag == "i":
            self._italic = True
        elif tag == "code":
            self._code = True
        elif tag == "table":
            self._in_table = True
            self._table_rows = []
        elif tag == "tr":
            self._current_row = []
        elif tag in ("td", "th"):
            self._current_cell = []

    def handle_endtag(self, tag: str) -> None:
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = getattr(self, "_heading_level", 1)
            # 清理首尾纯空白的 span（标题前后可能混入换行）
            spans = self._current_spans[:]
            while spans and not spans[0].text.strip():
                spans.pop(0)
            while spans and not spans[-1].text.strip():
                spans.pop()
            if any(s.text.strip() for s in spans):
                heading = self.doc.add_heading(level=min(level, 6))
                _spans_to_runs(heading, spans)
            self._current_spans = []
            if hasattr(self, "_heading_level"):
                del self._heading_level
        elif tag == "p":
            if self._current_spans and any(s.text.strip() for s in self._current_spans):
                para = self.doc.add_paragraph()
                if self._blockquote_depth > 0:
                    para.paragraph_format.left_indent = Pt(24 * self._blockquote_depth)
                _spans_to_runs(para, self._current_spans)
            self._current_spans = []
        elif tag in ("ul", "ol"):
            self._list_depth = max(0, self._list_depth - 1)
        elif tag == "li":
            if self._current_spans:
                style = "List Number" if self._in_ordered else "List Bullet"
                if self._list_depth > 1:
                    style = f"{'List Number' if self._in_ordered else 'List Bullet'} {min(self._list_depth, 3)}"
                para = self.doc.add_paragraph(style=style)
                _spans_to_runs(para, self._current_spans)
            self._current_spans = []
        elif tag == "blockquote":
            self._blockquote_depth = max(0, self._blockquote_depth - 1)
        elif tag in ("strong", "b"):
            self._bold = False
        elif tag in ("em", "i"):
            self._italic = False
        elif tag == "code":
            self._code = False
        elif tag in ("td", "th"):
            self._current_row.append(self._current_cell)
            self._current_cell = []
        elif tag == "tr":
            if self._current_row:
                self._table_rows.append(self._current_row)
            self._current_row = []
        elif tag == "table":
            self._flush_table()
            self._in_table = False

    def handle_data(self, data: str) -> None:
        text = _decode_entities(data)
        self._flush_inline_to_span(text)

    def _flush_table(self) -> None:
        """把收集的表格行写入 Word 表格。"""
        if not self._table_rows:
            return
        max_cols = max(len(row) for row in self._table_rows)
        table = self.doc.add_table(rows=len(self._table_rows), cols=max_cols)
        try:
            table.style = "Light Grid Accent 1"
        except KeyError:
            pass  # 样式不存在时用默认表格
        for r, row_cells in enumerate(self._table_rows):
            for c in range(max_cols):
                cell = table.rows[r].cells[c]
                cell_spans = row_cells[c] if c < len(row_cells) else []
                if cell_spans:
                    para = cell.paragraphs[0]
                    _spans_to_runs(para, cell_spans)
                    if r == 0:  # 表头加粗
                        for run in para.runs:
                            run.bold = True

    def error(self, message: str) -> None:  # pragma: no cover
        pass


def markdown_to_docx_bytes(markdown_text: str, title: str | None = None) -> bytes:
    """把会议纪要 Markdown 转成 docx 字节流。

    Args:
        markdown_text: Markdown 纪要正文
        title: 可选文档标题（用于元数据，正文里的 # 也作标题）

    Returns:
        docx 文件的字节内容，可直接写入文件或作为响应体返回。
    """
    document = Document()
    # 设置默认正文字体（中文兼容）
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # 解析 markdown → HTML
    # 不用 nl2br 扩展——它会把标题间的空行变成 <br>，导致标题前混入换行符。
    # markdown 库默认的段落处理已经足够。
    extensions = ["tables", "fenced_code"]
    html = md_lib.markdown(markdown_text, extensions=extensions)

    # 用解析器构建 docx
    builder = _DocxBuilder(document)
    builder.feed(html)
    builder.close()

    # 写入字节流
    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


def write_docx(markdown_text: str, output_path: Any, title: str | None = None) -> None:
    """转 docx 并写入文件路径。"""
    docx_bytes = markdown_to_docx_bytes(markdown_text, title)
    from pathlib import Path
    path = Path(output_path)
    path.write_bytes(docx_bytes)
